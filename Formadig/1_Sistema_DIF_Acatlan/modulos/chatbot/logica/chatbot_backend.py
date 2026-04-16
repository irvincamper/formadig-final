from flask import Blueprint, request, jsonify, send_file
import requests
import os
import io
from datetime import datetime

# ⚠️ IMPORTACIONES CRÍTICAS PARA EXPORTACIÓN DE REPORTES
# Estas son necesarias para que el chatbot pueda generar archivos Excel, PDF, Word
try:
    import pandas as pd
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment
except ImportError as e:
    print(f"❌ ERROR CRÍTICO: pandas/openpyxl no instalados. Necesarios para Excel. Error: {e}")
    print(f"   Ejecuta: pip install pandas openpyxl")
    raise

try:
    from fpdf import FPDF
except ImportError as e:
    print(f"❌ ERROR CRÍTICO: fpdf2 no instalado. Necesario para PDF. Error: {e}")
    print(f"   Ejecuta: pip install fpdf2")
    raise

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError as e:
    print(f"❌ ERROR CRÍTICO: python-docx no instalado. Necesario para Word. Error: {e}")
    print(f"   Ejecuta: pip install python-docx")
    raise

# =========================================================
# 🏗️ CONFIGURACIÓN Y CONEXIÓN
# =========================================================
try:
    from supabase import create_client, Client
except ImportError:
    print("CRÍTICO: Debes ejecutar 'pip install supabase' en tu terminal.")
    exit(1)

# Crear Blueprint para chatbot
chatbot_bp = Blueprint('chatbot', __name__, url_prefix='/api/chatbot')

# Credenciales Supabase (desde environment variables SOLAMENTE para producción)
SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_KEY = os.getenv("SUPABASE_KEY")
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
GEMINI_URL = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-3-flash-preview:generateContent?key={GEMINI_API_KEY}"


try:
    supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)
except Exception as e:
    print(f"ATENCIÓN: Error al conectar con Supabase: {e}")
    supabase = None

# =========================================================
# 🧠 LÓGICA DE OBTENCIÓN DE CONTEXTO (RAG LIGHT)
# =========================================================

# Esquemas de tablas para guía de carga masiva
ESQUEMAS_TABLAS = {
    "traslados": ["paciente_nombre", "paciente_curp", "destino_hospital", "fecha_viaje", "acompanante_nombre", "telefono"],
    "desayunos_eaeyd": ["nombres", "apellidos", "curp", "localidad", "tipo_apoyo", "escuela"],
    "hospitales": ["nombre_hospital", "ubicacion"]
}

# Configuración de Columnas para Reportes Profesionales
REPORT_CONFIG = {
    "traslados": {
        "columns": ["paciente_nombre", "destino_hospital", "fecha_viaje", "estatus"],
        "labels": ["PACIENTE", "HOSPITAL DESTINO", "FECHA VIAJE", "ESTATUS"],
        "widths": [80, 100, 40, 40] # Para PDF
    },
    "desayunos_eaeyd": {
        "columns": ["nombres", "apellidos", "localidad", "tipo_apoyo", "escuela"],
        "labels": ["NOMBRES", "APELLIDOS", "LOCALIDAD", "APOYO", "INSTITUCIÓN"],
        "widths": [60, 60, 50, 40, 60]
    }
}

def filtrar_datos_reporte(data, table_name):
    if table_name not in REPORT_CONFIG or not data:
        return data, [] # Por defecto enviamos todo si no hay config
    
    config = REPORT_CONFIG[table_name]
    filtered_data = []
    
    for row in data:
        new_row = {}
        for i, col in enumerate(config["columns"]):
            label = config["labels"][i]
            new_row[label] = row.get(col, '')
        filtered_data.append(new_row)
        
    return filtered_data, config["labels"]

def obtener_metricas_globales(tema=None):
    if not supabase: return "No hay conexión a la DB."
    
    resumen = "\n--- MÉTRICAS PARA REPORTES (DIF ACATLÁN) ---\n"
    try:
        if tema == "traslados" or tema is None:
            # Conteos de Traslados
            t_count = supabase.table('traslados').select("*", count="exact").limit(1).execute().count
            t_pend = supabase.table('traslados').select("*", count="exact").eq('estatus', 'PENDIENTE').limit(1).execute().count
            resumen += f"- TOTAL TRASLADOS: {t_count} ({t_pend} Pendientes)\n"
        
        if tema == "desayunos" or tema is None:
            # Conteos de Desayunos
            d_frios = 0; d_calientes = 0; d_eaeyd = 0
            try: d_frios = supabase.table('desayunos_frios').select("*", count="exact").limit(1).execute().count or 0
            except: pass
            try: d_calientes = supabase.table('desayunos_calientes').select("*", count="exact").limit(1).execute().count or 0
            except: pass
            try: d_eaeyd = supabase.table('desayunos_eaeyd').select("*", count="exact").limit(1).execute().count or 0
            except: pass
            
            d_total = d_frios + d_calientes + d_eaeyd
            resumen += f"- TOTAL BENEFICIARIOS DESAYUNOS: {d_total}\n"
            if d_frios > 0: resumen += f"  - Programa Desayunos Fríos: {d_frios}\n"
            if d_calientes > 0: resumen += f"  - Programa Desayunos Calientes: {d_calientes}\n"
            if d_eaeyd > 0: resumen += f"  - Otros Programas EAEyD: {d_eaeyd}\n"
        
        if tema is None:
            # Hospitales
            h_count = supabase.table('hospitales').select("*", count="exact").limit(1).execute().count
            resumen += f"- HOSPITALES REGISTRADOS EN CATÁLOGO: {h_count}\n"
        
    except Exception as e:
        resumen += f"Error al consolidar métricas reales: {e}\n"
        
    return resumen

def obtener_contexto_db(user_email=None, user_role=None, intencion=None):
    if not supabase: return ""

    contexto = "\n--- CONTEXTO DE LA BASE DE DATOS (DIF ACATLÁN) ---\n"
    tema = intencion.get('tema') if intencion else None
    estatus_filtro = intencion.get('estatus') if intencion else None
    
    # 1. Catálogo de Hospitales (Solo si es general o traslados)
    if tema in ["traslados", None]:
        try:
            hosp_res = supabase.table('hospitales').select('nombre_hospital').limit(20).execute()
            hospitales = [h['nombre_hospital'] for h in hosp_res.data]
            contexto += f"Hospitales registrados: {', '.join(hospitales)}\n"
        except: pass

    # 2. Conocimiento General (FAQ/Reglas)
    try:
        conoc_res = supabase.table('chatbot_conocimiento').select('titulo, contenido').limit(10).execute()
        for item in conoc_res.data:
            contexto += f"Regla/Info [{item['titulo']}]: {item['contenido']}\n"
    except: pass

    # 3. Datos Transaccionales (Filtrados por usuario e intención)
    if user_email:
        try:
            # Buscar perfil por email
            perfil_res = supabase.table('perfiles').select('id, rol').eq('nombre_usuario', user_email).execute()
            if perfil_res.data:
                user_id = perfil_res.data[0]['id']
                is_admin = perfil_res.data[0]['rol'] in ['admin', 'directora', 'desarrollador', 'admin_traslados', 'admin_desayunos']

                # Traslados (Si el tema es traslados o no se especificó)
                if tema in ["traslados", None]:
                    query_t = supabase.table('traslados').select('paciente_nombre, destino_hospital, fecha_viaje, estatus')
                    if not is_admin: query_t = query_t.eq('registrado_por', user_id)
                    if estatus_filtro: query_t = query_t.eq('estatus', estatus_filtro)
                    
                    t_res = query_t.limit(5).execute()
                    if t_res.data:
                        contexto += f"Traslados {'específicos ' if estatus_filtro else ''}encontrados:\n"
                        for t in t_res.data:
                            contexto += f"- Paciente: {t['paciente_nombre']}, Destino: {t['destino_hospital']}, Fecha: {t['fecha_viaje']}, Estatus: {t['estatus']}\n"
                
                # Desayunos
                if tema in ["desayunos", None]:
                    registros_d = []
                    tablas_d = ['desayunos_frios', 'desayunos_calientes', 'desayunos_eaeyd']
                    if intencion and intencion.get('sub_tema'):
                        tablas_d = [t for t in tablas_d if intencion['sub_tema'] in t] or tablas_d

                    for t in tablas_d:
                        try:
                            # QUERY CORREGIDA: Selección de columnas explícita
                            columnas = "nombres, apellidos, localidad, estatus"
                            q = supabase.table(t).select(columnas)
                            
                            if not is_admin: q = q.eq('registrado_por', user_id)
                            if estatus_filtro: q = q.eq('estatus', estatus_filtro)
                            
                            res = q.limit(5).execute()
                            if res.data: 
                                for r in res.data:
                                    r['origen'] = t.replace('desayunos_', '').title()
                                    registros_d.append(r)
                        except Exception as e:
                            print(f"Error consultando tabla {t}: {e}")
                            continue

                    if registros_d:
                        contexto += f"Registros de desayunos {'específicos ' if estatus_filtro else ''}encontrados:\n"
                        for d in registros_d[:10]:
                            contexto += f"- [{d['origen']}] {d['nombres']} {d['apellidos']} ({d['localidad']}) - Estatus: {d.get('estatus', 'N/A')}\n"
        except Exception as e:
            print(f"Error cargando contexto transaccional: {e}")

    return contexto

def detectar_intencion(mensaje):
    msg = mensaje.lower()
    intenciones = {
        "tema": None,
        "sub_tema": None,
        "estatus": None,
        "es_general": False
    }
    
    # 1. Detectar Tema Principal
    if any(k in msg for k in ['traslado', 'paciente', 'viaje', 'hospital']):
        intenciones["tema"] = "traslados"
    elif any(k in msg for k in ['desayuno', 'frio', 'caliente', 'eaeyd', 'apoyo', 'comida', 'beneficiario']):
        intenciones["tema"] = "desayunos"
        if 'frio' in msg: intenciones["sub_tema"] = "frios"
        elif 'caliente' in msg: intenciones["sub_tema"] = "calientes"
        elif 'eaeyd' in msg: intenciones["sub_tema"] = "eaeyd"

    # 2. Detectar Estatus específico (Coincidir con CamelCase de la DB)
    if any(k in msg for k in ['rechazado', 'rechazada', 'denegado', 'cancelado']):
        intenciones["estatus"] = "Rechazado"
    elif any(k in msg for k in ['pendiente', 'espera', 'proceso']):
        intenciones["estatus"] = "Lista de Espera"
    elif any(k in msg for k in ['aprobado', 'aceptado', 'autorizado', 'activo']):
        intenciones["estatus"] = "Aprobado"
    elif any(k in msg for k in ['entregado', 'completado', 'finalizado']):
        intenciones["estatus"] = "Aprobado" # O ajuste según tabla

    # 3. Detectar si pide formato CSV
    if '.csv' in msg or ' csv' in msg:
        intenciones["pide_csv"] = True

    # 4. Detectar si es una pregunta muy general o fuera de lugar
    keywords_sistema = ['traslado', 'desayuno', 'hospital', 'apoyo', 'curp', 'solicitud', 'sistema', 'formadig', 'dif', 'acatlan', 'reporte', 'cargar', 'excel', 'hola', 'buenos días', 'buenas tardes']
    if not any(k in msg for k in keywords_sistema):
        intenciones["es_general"] = True

    return intenciones


# 📥 LÓGICA DE EXPORTACIÓN (EXCEL / WORD)
# =========================================================
def crear_excel_reporte(data, titulo):
    output = io.BytesIO()
    df = pd.DataFrame(data)
    
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df.to_excel(writer, index=False, sheet_name='Reporte')
        workbook = writer.book
        worksheet = writer.sheets['Reporte']
        
        # Formato de cabecera
        header_format = workbook.add_format({
            'bold': True,
            'bg_color': '#005c53',
            'font_color': 'white',
            'border': 1,
            'align': 'center'
        })
        
        for col_num, value in enumerate(df.columns.values):
            worksheet.write(0, col_num, value, header_format)
            worksheet.set_column(col_num, col_num, 25) # Ancho más generoso
            
    output.seek(0)
    return output

def crear_word_reporte(data, titulo):
    doc = Document()
    
    # Título Institucional
    header = doc.add_heading('SISTEMA FORMADIG - DIF ACATLÁN', 0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"Reporte Generado: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    doc.add_heading(titulo, level=1)
    
    if not data:
        doc.add_paragraph("No se encontraron registros.")
    else:
        df = pd.DataFrame(data)
        columns = df.columns.tolist()
        table = doc.add_table(rows=1, cols=len(columns))
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(columns):
            hdr_cells[i].text = col
            
        for row in data:
            row_cells = table.add_row().cells
            for i, col in enumerate(columns):
                row_cells[i].text = str(row.get(col, ''))
                
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output

def crear_pdf_reporte(data, table_name, titulo):
    class PDF(FPDF):
        def header(self):
            # Banner Verde Superior
            self.set_fill_color(0, 92, 83)
            self.rect(0, 0, 300, 30, 'F')
            self.set_text_color(255, 255, 255)
            self.set_font('helvetica', 'B', 20)
            self.cell(0, 10, 'DIF MUNICIPAL ACATLÁN DE JUÁREZ', border=0, align='C', ln=1)
            self.set_font('helvetica', 'I', 12)
            self.cell(0, 10, 'Sistema de Control de Traslados y Programas (FORMADIG)', border=0, align='C', ln=1)
            self.ln(12)

        def footer(self):
            self.set_y(-20)
            self.set_font('helvetica', 'I', 8)
            self.set_text_color(100, 100, 100)
            self.cell(0, 10, f'Generado el {datetime.now().strftime("%d/%m/%Y")} - Página {self.page_no()}/{{nb}}', align='C')

    pdf = PDF(orientation='L', unit='mm', format='A4')
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()
    
    pdf.set_text_color(0, 0, 0)
    pdf.set_font('helvetica', 'B', 16)
    pdf.cell(0, 10, titulo, ln=1, align='L')
    pdf.ln(5)
    
    if not data:
        pdf.cell(0, 10, 'Sin registros.', ln=1)
    else:
        config = REPORT_CONFIG.get(table_name, {"widths": [50]*len(data[0].keys())})
        labels = list(data[0].keys())
        widths = config["widths"]

        # Cabecera
        pdf.set_font('helvetica', 'B', 11)
        pdf.set_fill_color(0, 92, 83)
        pdf.set_text_color(255, 255, 255)
        for i, label in enumerate(labels):
            pdf.cell(widths[i], 12, label, border=1, align='C', fill=True)
        pdf.ln()

        # Filas
        pdf.set_font('helvetica', '', 10)
        pdf.set_text_color(0, 0, 0)
        for row in data:
            # Calcular altura necesaria según el texto más largo en la fila
            row_height = 10
            for i, label in enumerate(labels):
                val = str(row.get(label, ''))
                # Simular multicelda o truncar para mantener diseño
                pdf.cell(widths[i], row_height, val[:35], border=1, align='L')
            pdf.ln()

    output = io.BytesIO()
    pdf.output(output)
    output.seek(0)
    return output

@chatbot_bp.route('/export', methods=['GET'], strict_slashes=False)
def export_report():
    table_name = request.args.get('table', 'traslados')
    file_format = request.args.get('format', 'excel')
    
    if not supabase:
        return "Error: No hay conexión a la base de datos.", 500
        
    try:
        raw_data = []
        if table_name == 'desayunos_eaeyd':
            # AGREGACIÓN INTELIGENTE PARA REPORTE UNIFICADO
            for t in ['desayunos_frios', 'desayunos_calientes', 'desayunos_eaeyd']:
                try:
                    res = supabase.table(t).select("*").execute()
                    if res.data:
                        for r in res.data:
                            # Asegurar columna tipo_apoyo si no existe
                            if 'tipo_apoyo' not in r:
                                r['tipo_apoyo'] = t.replace('desayunos_', '').title()
                            raw_data.append(r)
                except: continue
        else:
            res = supabase.table(table_name).select("*").execute()
            raw_data = res.data if res.data else []
        
        # 2. Filtrar y Profesionalizar
        data, _ = filtrar_datos_reporte(raw_data, table_name)
        
        filename = f"Reporte_{table_name}_{datetime.now().strftime('%Y%m%d')}"
        
        if file_format == 'excel':
            file_data = crear_excel_reporte(data, f"Reporte de {table_name}")
            return send_file(file_data, as_attachment=True, download_name=f"{filename}.xlsx", mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        elif file_format == 'pdf':
            file_data = crear_pdf_reporte(data, table_name, f"REPORTE OFICIAL: {table_name.upper()}")
            return send_file(file_data, as_attachment=True, download_name=f"{filename}.pdf", mimetype='application/pdf')
        else:
            file_data = crear_word_reporte(data, f"Reporte Institucional: {table_name.title()}")
            return send_file(file_data, as_attachment=True, download_name=f"{filename}.docx", mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    except Exception as e:
        print(f"Error en exportación: {e}")
        return f"Error al generar el archivo: {str(e)}", 500

# =========================================================
# 📤 LOGICA DE CARGA MASIVA (UPLOAD) - V2 (MAPEADO INTELIGENTE)
# =========================================================

# Mapeo de Etiquetas Profesionales -> Columnas de BD
REVERSE_MAPPING = {
    "paciente": "paciente_nombre",
    "hospital_destino": "destino_hospital",
    "fecha_viaje": "fecha_viaje",
    "estatus": "estatus",
    "nombres": "nombres",
    "apellidos": "apellidos",
    "localidad": "localidad",
    "apoyo": "tipo_apoyo",
    "institución": "escuela",
    "institucion": "escuela",
    "hospital": "destino_hospital"
}

@chatbot_bp.route('/upload', methods=['POST'], strict_slashes=False)
def upload_file():
    if not supabase:
        return jsonify({"error": "No hay conexión a la base de datos."}), 500
        
    if 'file' not in request.files:
        return jsonify({"error": "No se subió ningún archivo"}), 400
        
    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "Archivo vacío"}), 400

    try:
        # 1. Leer archivo según extensión
        ext = file.filename.split('.')[-1].lower()
        if ext == 'csv':
            df = pd.read_csv(file)
        elif ext in ['xlsx', 'xls']:
            df = pd.read_excel(file)
        else:
            return jsonify({"error": "Formato no soportado (usa CSV o Excel)"}), 400

        # Normalizar nombres de columnas (quitar espacios, minúsculas, tildes básicas)
        df.columns = [str(c).strip().lower().replace(' ', '_') for c in df.columns]
        
        # Aplicar Mapeo Inverso (Etiquetas -> BD)
        mapped_columns = {}
        for col in df.columns:
            if col in REVERSE_MAPPING:
                mapped_columns[col] = REVERSE_MAPPING[col]
        
        if mapped_columns:
            df.rename(columns=mapped_columns, inplace=True)
            
        cols = set(df.columns)

        # 2. Identificar Tabla Destino
        target_table = None
        if 'paciente_nombre' in cols or 'destino_hospital' in cols:
            target_table = 'traslados'
        elif 'escuela' in cols or 'tipo_apoyo' in cols or 'nombres' in cols:
            target_table = 'desayunos_eaeyd'
        else:
            return jsonify({
                "error": "No se pudo identificar el tipo de datos. Asegúrate de que el archivo tenga columnas como 'Paciente', 'Hospital', 'Nombres' o 'Apoyo'.",
                "detected_columns": list(cols)
            }), 400

        # 3. Limpiar y Preparar Datos
        # Filtrar solo columnas que existen en la tabla destino (según ESQUEMAS_TABLAS)
        if target_table in ESQUEMAS_TABLAS:
            valid_cols = ESQUEMAS_TABLAS[target_table]
            # Mantener solo las que están en el DF
            cols_to_keep = [c for c in df.columns if c in valid_cols]
            df = df[cols_to_keep]

        data = df.to_dict(orient='records')
        
        if not data:
            return jsonify({"error": "El archivo no contiene datos válidos para procesar."}), 400

        # Inserción masiva en Supabase
        res = supabase.table(target_table).insert(data).execute()
        
        return jsonify({
            "message": f"¡Carga Exitosa! Se importaron {len(data)} registros en {target_table.replace('_', ' ').title()}.",
            "count": len(data),
            "table": target_table
        })

    except Exception as e:
        print(f"Error en Carga Masiva: {e}")
        return jsonify({"error": f"Error al procesar el archivo: {str(e)}"}), 500

# =========================================================
# 🛡️ MOTOR DE CONTINGENCIA (FALLBACK LOCAL)
# =========================================================
def generar_respuesta_local(mensaje, email, role):
    mensaje = mensaje.lower()
    metricas = obtener_metricas_globales()
    
    # Detección de Temas
    es_reporte = 'reporte' in mensaje or 'resumen' in mensaje or 'estadistica' in mensaje or 'cuanto' in mensaje
    es_traslado = 'traslado' in mensaje or 'paciente' in mensaje or 'viaje' in mensaje
    es_desayuno = 'desayuno' in mensaje or 'frio' in mensaje or 'caliente' in mensaje or 'comida' in mensaje
    
    respuesta = "AVISO INSTITUCIONAL: El servicio de Inteligencia Artificial está temporalmente fuera de línea.\n"
    respuesta += "Como asistente del sistema FORMADIG del DIF Acatlán, puedo proporcionarle la siguiente información técnica basada en los registros actuales:\n\n"
    
    if es_reporte or es_traslado or es_desayuno:
        respuesta += "Métricas consolidadas del sistema:\n"
        respuesta += metricas + "\n"
        
        # Ofrecer descargas si pidió reporte
        if es_reporte or es_traslado:
            respuesta += "\nGestión de archivos para Traslados Médicos:\n"
            respuesta += "EXPORT_BUTTONS:\n"
            respuesta += f"- [Descargar Reporte (Excel)](/api/chatbot/export?table=traslados&format=excel)\n"
            respuesta += f"- [Descargar Reporte (PDF)](/api/chatbot/export?table=traslados&format=pdf)\n"
        
        if es_reporte or es_desayuno:
            respuesta += "\nGestión de archivos para Programas Alimentarios:\n"
            respuesta += "EXPORT_BUTTONS:\n"
            respuesta += f"- [Descargar Reporte (Excel)](/api/chatbot/export?table=desayunos_eaeyd&format=excel)\n"
            respuesta += f"- [Descargar Reporte (PDF)](/api/chatbot/export?table=desayunos_eaeyd&format=pdf)\n"
    else:
        respuesta += "Para asistirle adecuadamente en este modo de contingencia, por favor solicite información sobre 'reportes', 'traslados' o 'programas alimentarios'."
        
    return respuesta


# =========================================================
# 📡 RUTA PRINCIPAL
# =========================================================
@chatbot_bp.route('/ask', methods=['POST'], strict_slashes=False)
def ask_gemini():
    data = request.json
    user_message = data.get('message', '')
    user_email = data.get('email', '')
    user_role = data.get('role', '')

    if not user_message:
        return jsonify({"error": "No message provided"}), 400

    # 1. Detectar intención y filtrar contexto
    intencion = detectar_intencion(user_message)
    contexto_db = obtener_contexto_db(user_email, user_role, intencion)
    metricas = obtener_metricas_globales(intencion.get('tema'))

    mensaje_predeterminado = ""
    if intencion.get('es_general'):
        mensaje_predeterminado = """
        NOTA: El usuario parece estar haciendo una pregunta general o fuera del contexto del DIF Acatlán. 
        Si la pregunta no tiene NADA que ver con el DIF, FORMADIG, traslados, desayunos o gestión municipal, 
        debes responder cortésmente: 'Lo siento, como asistente del DIF Acatlán, solo puedo ayudarte con temas relacionados a nuestros programas sociales y al sistema FORMADIG. ¿En qué más puedo apoyarte respecto a estas áreas?'
        """

    nota_csv = ""
    if intencion.get('pide_csv'):
        nota_csv = "ADVERTENCIA: El usuario ha solicitado un archivo CSV. Debes responder explícitamente: 'No genero archivos en formato CSV, ya que no es un estándar institucional para reportes finales. Sin embargo, puedo generar el archivo en Excel, Word o PDF para su correcta visualización y manejo. ¿En qué más puedo ayudarle?'"

    prompt_sistema = f"""
    PERSONALIDAD Y ROL:
    Eres el Asistente Virtual oficial del DIF Municipal de Acatlán de Juárez, diseñado exclusivamente para el sistema FORMADIG. 
    Tu función primordial es el procesamiento, organización y generación de reportes automáticos de todos los módulos del sistema.

    ÁREAS DE COMPETENCIA:
    1. Traslados Médicos: Seguimiento y reportes sobre traslados a instituciones de salud.
    2. Programas Alimentarios (EAEyD): Información de beneficiarios de Desayunos Fríos y Calientes.
    3. Catálogo de Hospitales: Consulta de centros hospitalarios registrados.
    4. Gestión de Datos: Generación de reportes institucionales en formatos Excel, Word y PDF.

    FORMATO DE RESPUESTA (SALUDO INICIAL):
    Si el usuario te saluda (ej. 'hola', 'buenos días'), DEBES responder con una presentación institucional completa que incluya tus funciones principales, similar a:
    'Buenos días. Soy el asistente virtual del DIF Municipal de Acatlán (sistema FORMADIG). Estoy a su disposición para asistirle en la gestión y consulta de información relacionada con los programas de esta institución. Mis funciones principales incluyen...' (y listas las 4 áreas).

    CONTEXTO TÉCNICO REGISTRADO:
    {contexto_db}
    {metricas}
    {mensaje_predeterminado}
    {nota_csv}

    REGLAS DE FORMATO Y EXPORTACIÓN:
    1. GENERACIÓN DE REPORTES: Presenta datos en tablas markdown o listas claras.
    2. DESCARGA DE ARCHIVOS: Si piden Excel, Word o PDF, incluye SIEMPRE el bloque EXPORT_BUTTONS.
       - TRASLADOS: table=traslados
       - DESAYUNOS: table=desayunos_eaeyd
       Ejemplo bloque:
       EXPORT_BUTTONS:
       - [Descargar en Excel](/api/chatbot/export?table=traslados&format=excel)
       - [Descargar en Word](/api/chatbot/export?table=traslados&format=word)
       - [Descargar en PDF](/api/chatbot/export?table=traslados&format=pdf)
    3. NO CSV: Si solicitan CSV, utiliza la advertencia proporcionada en el contexto técnico.

    INSTRUCCIONES DE ESTILO (CRÍTICO):
    1. Tono puramente FORMAL, INSTITUCIONAL y PROFESIONAL (estilo ministerial).
    2. PROHIBIDO el uso de emojis en cualquier parte de la respuesta.
    3. El lenguaje debe ser directo, experto y servicial.
    4. Si no hay datos (ej. no hay registros pendientes), informe de manera sobria: 'Tras realizar la consulta en el sistema FORMADIG, no se han localizado registros bajo los criterios especificados'.

    USUARIO AUTENTICADO: {user_email} (Rol: {user_role})
    SOLICITUD DEL USUARIO: {user_message}
    """

    payload = {
        "contents": [{
            "parts": [{ "text": prompt_sistema }]
        }]
    }

    try:
        response = requests.post(GEMINI_URL, json=payload, timeout=8) # Timeout preventivo
        response_data = response.json()
        
        candidates = response_data.get('candidates', [])
        if candidates:
            bot_response = candidates[0].get('content', {}).get('parts', [{}])[0].get('text', 'Lo siento, hubo un problema al generar la respuesta.')
        else:
            bot_response = generar_respuesta_local(user_message, user_email, user_role)
            
        return jsonify({"response": bot_response})
    except Exception as e:
        print(f"Error en Chatbot Backend (AI Falló): {e}")
        # Intentar respuesta local por fallas de servidor/red
        bot_response = generar_respuesta_local(user_message, user_email, user_role)
        return jsonify({"response": bot_response})


# ============================================================================
# NOTA: El blueprint 'chatbot_bp' se registra en la app maestra
# ============================================================================